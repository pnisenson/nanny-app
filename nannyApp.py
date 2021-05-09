#from flask import Flask, render_template, redirect, jsonify
from flask import Flask, render_template, redirect, request
import os
from secret import phoneNo

# Create an instance of Flask
app = Flask(__name__)

# To send app link each morning when app is called
def send_link(url):
	import ezgmail
	ezgmail.send(phoneNo,'', f'Fill form for nanny: {url}')

# Create public URL for localhost:5000 so phone can access undeployed app
def start_ngrok():
	from pyngrok import ngrok
	url = ngrok.connect(5000).public_url
	return url

send_link(start_ngrok())

# Route to render index.html template using data from Mongo
@app.route("/")
def home():
	return redirect("/index.html")

@app.route("/index.html", methods=['GET', 'POST'])
def main():
	#Return template and data
	errors= []
	if request.method == "POST":
		try:
			if os.path.exists("nannyinfo.docx"):
				os.remove("nannyinfo.docx")
			writeFile(request.form)
			# printFile()
			recordData(request.form)
			return "Thank you for completing the form"
		except:
			errors.append("Form not complete. Please make sure it's valid and try again.")
			return render_template('index.html', errors=errors)
	return render_template("index.html")

def writeFile(inputs):
	from docx import Document
	doc = Document()
	doc.add_paragraph(f"This morning, Joey woke up at {inputs['wakeUp']} AM.")
	doc.add_paragraph(f"Her first nap should be at {inputs['firstNap']} AM.")
	doc.add_paragraph(f"For lunch today, we have {inputs['lunch']}.")
	doc.add_paragraph(f"For dinner today, we have {inputs['snack']}.")
	doc.add_paragraph(f"As a reminder {inputs['reminder']}.")
	doc.save("nannyinfo.docx")

def printFile():
	import subprocess
	subprocess.Popen(["C:\\Program Files (x86)\\Microsoft Office\\root\\Office16\\WINWORD.EXE", "nannyinfo.docx", "/mFilePrintDefault", "/mFileExit", "/mFileCloseOrExit"]).communicate()

def recordData(inputs):
	from openpyxl import load_workbook
	from datetime import date
	new_row_data = [date.today().strftime("%m/%d/%Y")]
	for i in inputs.values():
		new_row_data.append(i)
	wb = load_workbook("weeklyinput.xlsx")
	# Select First Worksheet
	ws = wb.worksheets[0]
	ws.append(new_row_data)
	wb.save("weeklyinput.xlsx")



if __name__ == "__main__":
    app.run(debug=True)
